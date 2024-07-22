import io
import csv
from fastapi import FastAPI, UploadFile, File, HTTPException
from typing import List
from fastapi import FastAPI, HTTPException
import re
from fastapi import FastAPI, File, UploadFile, HTTPException
from typing import List
from io import BytesIO
from pypdf import PdfReader
import pandas as pd
import re
import google.generativeai as genai
from pydantic import BaseModel
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_community.utilities import SQLDatabase
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI
from langchain_groq import ChatGroq
from langchain_anthropic import ChatAnthropic
import pandas as pd
import re
from PyPDF2 import PdfReader
from io import BytesIO
import pytesseract
from PIL import Image
import io
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, File, UploadFile, HTTPException, Body
from fastapi.responses import JSONResponse
from fastapi import FastAPI, File, UploadFile, Form
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from typing import List
import os
import docx
from typing import Optional
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
# import textract
import textwrap
import requests
from bs4 import BeautifulSoup
import PIL.Image
from io import BytesIO
from PIL import Image
from pydantic import BaseModel
from openpyxl import load_workbook
from fastapi.responses import HTMLResponse
from anthropic import Anthropic

load_dotenv()

app = FastAPI()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# read all pdf/csv/doc/docx/txt files and return text
def get_file_text(file_contents, file_extension):
    text = ""
    try:
        if file_extension == "pdf":
            pdf_reader = PdfReader(io.BytesIO(file_contents))
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file_extension == "csv":
            csv_reader = csv.reader(io.StringIO(file_contents.decode("utf-8")))
            for row in csv_reader:
                text += ' '.join(row) + '\n'
        elif file_extension == "docx":
            doc = docx.Document(io.BytesIO(file_contents))
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
        elif file_extension == "doc":
            doc = docx.Document(io.BytesIO(file_contents))
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
        elif file_extension == "txt":
            text = file_contents.decode("utf-8")
        elif file_extension == "xlsx":
            workbook = load_workbook(io.BytesIO(file_contents), data_only=True)
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += ' '.join([str(cell) if cell is not None else '' for cell in row]) + '\n'
    except Exception as e:
        print(f"Error processing file: {e}")
    return text


# website scrapping 
def get_website_text(url):
    try:
        response = requests.get(url)
        response.raise_for_status() 
        soup = BeautifulSoup(response.text, "html.parser")
        text = soup.get_text()
        return text
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Error fetching website: {str(e)}")
    
@app.post("/upload_website_url")
async def upload_website(url: str):
    try:
        text = get_website_text(url)
        chunks = get_text_chunks(text)
        get_vector_store(chunks)
        return JSONResponse(content={"message": "Website uploaded and processed successfully."}, status_code=200)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# split text into chunks
def get_text_chunks(text):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=1000)
    chunks = splitter.split_text(text)
    return chunks  

# get embeddings for each chunk
def get_vector_store(chunks):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001")  
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

    
def get_conversational_chain():
    prompt_template = """
    Provide a thorough and insightful answer to the question based on the given context. If the answer is not directly available in the context, use your knowledge and reasoning abilities to infer and extrapolate relevant information to formulate a comprehensive response. Your goal is to provide as much useful information as possible, even if it requires making logical connections or drawing reasonable conclusions beyond the literal context.And also provide default suggestion

    Context:
    {context}

    Question:
    {question}

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro",
                             temperature=0.3)

    prompt = PromptTemplate(template = prompt_template, input_variables = ["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain

# New Hist
history = []

def user_input(user_question):
    user_question = user_question
    
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001")  # type: ignore

    new_db = FAISS.load_local("faiss_index", embeddings)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()

    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True,)

        
    # Store the question and answer in the history
    history.append({"question": user_question, "answer": response})
        
    return response

# website url chatting api
@app.post("/upload_pdf_txt_csv_doc_docx")
async def upload_files(files: List[UploadFile] = File(...)):
    try:
        text = ""
        for uploaded_file in files:
            contents = await uploaded_file.read()
            filename, file_extension = os.path.splitext(uploaded_file.filename)
            text += get_file_text(contents, file_extension[1:])
        chunks = get_text_chunks(text)
        get_vector_store(chunks)
        return JSONResponse(content={"message": "Files uploaded and processed successfully."}, status_code=200)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/chat_with_uploaded_files")
async def chat_with_user(user_question: str):
    try:
        response = user_input(user_question)
        return JSONResponse(content=response, status_code=200)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
      

# Gemini Chatting
# Configuration
genai.configure(api_key=os.getenv("GOOGLE_API_KEY")) 
model = genai.GenerativeModel("gemini-pro")
chat = model.start_chat(history=[])

# Helper function
def get_gemini_response(question):
    response = chat.send_message(question, stream=True)
    return response

@app.post("/chat_with_google_gemini")
async def generate_content(prompt: str):
    try:
        prompt_parts = [prompt]
        response = model.generate_content(prompt_parts)
        response.resolve()
        output = []
        for chunk in response:
            output.append(chunk.text)
        # Store prompt and response in history
        history.append({"prompt": prompt, "response": output})
        return {"output": output}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
     
# Vision API'S
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

def to_markdown(text):
    text = text.replace('*', '\*')
    return textwrap.indent(text, '>', predicate=lambda _: True)

# image uploading 
@app.post("/upload_image_and_write_prompt")
def vision_upload(prompt: str, img_file: UploadFile = File(...)):
    try:
        img = PIL.Image.open(img_file.file)
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content([prompt, img])
        markdown_text = to_markdown(response.text)
        history.append({"prompt": prompt, "response": markdown_text})
        return {"output": markdown_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# website url image uploading
@app.post("/upload_image_url_write_prompt")
def vision_url(prompt: str, img_url: str):
    try:
        response = requests.get(img_url, stream=True)
        response.raise_for_status()
        
        # Check if the response contains image data
        if 'image' not in response.headers['Content-Type']:
            raise HTTPException(status_code=400, detail="The provided URL does not point to an image")

        img = PIL.Image.open(io.BytesIO(response.content))
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content([prompt, img])
        markdown_text = to_markdown(response.text)
        history.append({"prompt": prompt, "response": markdown_text})
        return {"output": markdown_text}

    # return markdown_text
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

    
# Function to load OpenAI model and get response
def get_gemini_response(input_text, image_data, prompt):
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input_text, image_data[0], prompt])
    return response.text

# Function to process uploaded image
def process_uploaded_image(uploaded_file):
    if uploaded_file is not None:
        image = Image.open(io.BytesIO(uploaded_file.file.read()))
        return image
    else:
        raise FileNotFoundError("No file uploaded")

# History Maintaining
@app.get("/history_of_response")
def get_history():
    try:
        return JSONResponse(content=history ,status_code=200)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) 
    

# Chatting with DataBase    
class DBConnectionInfo(BaseModel):
    host: str = "localhost"
    port: str = "5432"
    user: str = "postgres"
    password: str = "Nazar123"
    database: str = "godaam"         

db = None

def init_database(user: str, password: str, host: str, port: str, database: str) -> SQLDatabase:
    db_uri = f"postgresql://{user}:{password}@{host}:{port}/{database}"
    return SQLDatabase.from_uri(db_uri)

def get_sql_chain(db):
    template = """
        You are a data analyst at a company. You are interacting with a user who is asking you questions about the company's database.
        Based on the table schema below, write a SQL query that would answer the user's question. Take the conversation history into account.
        Never show a password if user ask you ok tell them it is privacy issue.
        
        <SCHEMA>{schema}</SCHEMA>
        
        
        Write only the SQL query and nothing else. Do not wrap the SQL query in any other text, not even backticks.
        
        For example:
        Question: which 3 artists have the most tracks?
        SQL Query: SELECT ArtistId, COUNT(*) as track_count FROM Track GROUP BY ArtistId ORDER BY track_count DESC LIMIT 3;
        Question: Name 10 artists
        SQL Query: SELECT Name FROM Artist LIMIT 10;
        
        Your turn:
        
        Question: {question}
        SQL Query:
    """
    
    prompt = ChatPromptTemplate.from_template(template)
    
    # llm = ChatOpenAI(model="gpt-4-0125-preview")
    # llm = ChatGroq(model="mixtral-8x7b-32768", temperature=0.3)
    # llm = ChatGoogleGenerativeAI(model="gemini-pro",temperature=0.3)
    llm = ChatAnthropic(model="claude-3-opus-20240229", anthropic_api_key=os.getenv("ANTHROPIC_API_KEY"))
    # llm = ChatAnthropic(model="claude-3.5-sonnet", anthropic_api_key=os.getenv("ANTHROPIC_API_KEY"))
  
    def get_schema(_):
        return db.get_table_info()
  
    return (
        RunnablePassthrough.assign(schema=get_schema)
        | prompt
        | llm
        | StrOutputParser()
    )

@app.post("/connect_to_postgresql")
async def connect_to_database(info: DBConnectionInfo):
    global db
    try:
        db = init_database(info.user, info.password, info.host, info.port, info.database)
        return {"message": "Successfully Connected to database!"}
    except Exception as e:
        error_message = f"Failed to connect to the database: {str(e)}"
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/chat_with_postgresql_database")
async def send_user_query(query: str):
    global db
    if db is None:
        raise HTTPException(status_code=400, detail="Database connection not established")
    try:
        sql_chain = get_sql_chain(db)
        response = get_response(query, db)
        history.append({"query": query, "response": response})
        return {"response": response}
    except Exception as e:
        # Handle any exceptions that may occur during the execution of the code
        error_message = f"An error occurred: {str(e)}"
        raise HTTPException(status_code=500, detail=error_message)

def get_response(user_query: str, db: SQLDatabase):
    sql_chain = get_sql_chain(db)

    template = """
    You are a data analyst at a company. You are interacting with a user who is asking you questions about the company's database.
    Based on the table schema below, question, sql query, and sql response, write a natural language response.
    <SCHEMA>{schema}</SCHEMA>

    SQL Query: <SQL>{query}</SQL>
    User question: {question}
    SQL Response: {response}"""

    prompt = ChatPromptTemplate.from_template(template)

    # llm = ChatOpenAI(model="gpt-4-0125-preview")
    # llm = ChatGroq(model="mixtral-8x7b-32768", temperature=0.3)
    # llm = ChatGoogleGenerativeAI(model="gemini-pro",temperature=0.3)
    llm = ChatAnthropic(model="claude-3-opus-20240229", anthropic_api_key=os.getenv("ANTHROPIC_API_KEY"))
    # llm = ChatAnthropic(model="claude-3.5-sonnet", anthropic_api_key=os.getenv("ANTHROPIC_API_KEY"))


    chain = (
        RunnablePassthrough.assign(query=sql_chain).assign(
            schema=lambda _: db.get_table_info(),
            response=lambda vars: db.run(vars["query"]),
        )
        | prompt
        | llm
        | StrOutputParser()
    )

    return chain.invoke({
        "question": user_query,
    })

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)